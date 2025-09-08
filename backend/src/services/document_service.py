"""
文档处理服务
处理文档导出功能
"""

import os
import tempfile
from docx import Document
from docx.shared import Inches
import html2text
import re
import difflib
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class DocumentService:
    def __init__(self):
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = True
        self.html_converter.ignore_images = True

    def render_report_html(self, content, result, meta=None):
        """
        生成可打印的审校报告 HTML（用于前端预览或前端导出为 PDF）。
        - 原文为红色
        - 建议为默认颜色
        - 仅差异字符添加下划线
        """
        meta = meta or {}
        title = meta.get('title') or '审校报告'
        author = meta.get('author') or ''
        rules_mode = (meta.get('rules_mode') or '').upper() or 'LITE'
        qwen_enabled = bool(meta.get('qwen', True))
        stats = (result or {}).get('statistics') or {}
        issues = (result or {}).get('issues') or []

        def esc(s):
            return (s or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        def context_snippet(pos):
            try:
                s = pos.get('start', 0); e = pos.get('end', 0)
                left = max(0, s - 20); right = min(len(content), e + 20)
                snippet = content[left:right]
                return esc(snippet)
            except Exception:
                return ''

        # 简易 diff，返回两个带 <u> 的 HTML 片段
        def diff_html(orig, sug, is_original):
            orig = orig or ''
            sug = sug or ''
            sm = difflib.SequenceMatcher(a=orig, b=sug)
            parts = []
            for tag, i1, i2, j1, j2 in sm.get_opcodes():
                if is_original:
                    seg = esc(orig[i1:i2])
                    if tag in ('replace', 'delete') and seg:
                        parts.append(f'<u>{seg}</u>')
                    else:
                        parts.append(seg)
                else:
                    seg = esc(sug[j1:j2])
                    if tag in ('replace', 'insert') and seg:
                        parts.append(f'<u>{seg}</u>')
                    else:
                        parts.append(seg)
            return ''.join(parts)

        css = """
        <style>
          body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, 'Noto Sans CJK SC', 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif; margin: 20px; }
          h1, h2, h3 { margin: 0.6em 0; }
          .muted { color: #666; }
          .cover { border: 1px solid #eee; padding: 16px; border-radius: 8px; }
          .badge { display:inline-block; background:#f5f5f5; border:1px solid #e5e5e5; padding:2px 6px; border-radius:6px; font-size:12px; color:#333; }
          table { width: 100%; border-collapse: collapse; margin-top: 12px; }
          th, td { border: 1px solid #ddd; padding: 8px; vertical-align: top; }
          tr:nth-child(even) { background: #fafafa; }
          .sev-high { color: #d32f2f; font-weight: 600; }
          .sev-medium { color: #ef6c00; font-weight: 600; }
          .sev-warning { color: #616161; }
          .orig { color: #b91c1c; }
          .arrow { color: #9e9e9e; padding: 0 6px; }
          u { text-underline-offset: 4px; }
        </style>
        """

        rows = []
        for idx, it in enumerate(issues, start=1):
            t = esc(it.get('type') or '-')
            sev = (it.get('severity') or 'warning').lower()
            sev_class = {
                'high': 'sev-high', 'medium': 'sev-medium', 'warning': 'sev-warning'
            }.get(sev, 'sev-warning')
            pos = it.get('position') or {}
            pos_text = f"{pos.get('start', 0)}-{pos.get('end', 0)}"
            orig = it.get('original') or ''
            sug = it.get('suggestion') or ''
            msg = esc(it.get('message') or it.get('description') or '')
            ctx = context_snippet(pos)
            orig_html = diff_html(orig, sug, True)
            sug_html = diff_html(orig, sug, False)
            rows.append(f"""
              <tr>
                <td>{idx}</td>
                <td>{t}</td>
                <td class='{sev_class}'>{esc(it.get('severity') or 'warning')}</td>
                <td>{pos_text}</td>
                <td class='orig'>{orig_html}</td>
                <td>{sug_html}</td>
                <td>{msg}</td>
                <td><code>{ctx}</code></td>
              </tr>
            """)

        html = f"""
        <html><head><meta charset='utf-8' />{css}</head>
        <body>
          <div class='cover'>
            <h1>{esc(title)}</h1>
            <div class='muted'>作者：{esc(author)}　·　导出时间：{meta.get('export_time') or ''}</div>
            <div style='margin-top:8px;'>
              <span class='badge'>规则模式：{esc(rules_mode)}</span>
              <span class='badge'>LLM：{'启用' if qwen_enabled else '关闭'}</span>
              <span class='badge'>字数：{len(content)}</span>
              <span class='badge'>问题数：{stats.get('total_issues', 0)}</span>
            </div>
          </div>
          <h2>摘要统计</h2>
          <div class='muted'>错别字：{stats.get('typos',0)}　语法：{stats.get('grammar',0)}　标点：{stats.get('punctuation',0)}　敏感：{stats.get('sensitive',0)}</div>
          <h2>问题清单</h2>
          <table>
            <thead>
              <tr>
                <th>#</th><th>类型</th><th>严重度</th><th>位置</th><th>原文</th><th>建议</th><th>说明</th><th>上下文</th>
              </tr>
            </thead>
            <tbody>
              {''.join(rows)}
            </tbody>
          </table>
        </body></html>
        """
        return html

    def docx_from_report(self, content, result, title="审校报告", author="", meta=None):
        """
        纯 python-docx 生成结构化审校报告（方案B）。
        - 封面、摘要、问题表格
        - 原文：红色，差异片段下划线
        - 建议：默认色，差异片段下划线
        返回：bytes
        """
        meta = meta or {}
        doc = Document()

        # 封面
        if title:
            h = doc.add_heading(title, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(f"作者：{author or ''}    规则模式：{(meta.get('rules_mode') or '').upper() or 'LITE'}    LLM：{'启用' if meta.get('qwen', True) else '关闭'}")
        meta_run.font.size = Pt(10)

        # 摘要
        stats = (result or {}).get('statistics') or {}
        doc.add_heading('摘要统计', level=2)
        sum_p = doc.add_paragraph()
        sum_p.add_run(f"总问题数：{stats.get('total_issues', 0)}  错别字：{stats.get('typos',0)}  语法：{stats.get('grammar',0)}  标点：{stats.get('punctuation',0)}  敏感：{stats.get('sensitive',0)}")

        # 详细问题
        issues = (result or {}).get('issues') or []
        doc.add_heading('问题清单', level=2)
        cols = ['#','类型','严重度','位置','原文','建议','说明','上下文']
        table = doc.add_table(rows=1, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr_cells[i].text = c

        for idx, it in enumerate(issues, start=1):
            row = table.add_row().cells
            t = it.get('type') or '-'
            sev = it.get('severity') or 'warning'
            pos = it.get('position') or {}
            pos_text = f"{pos.get('start',0)}-{pos.get('end',0)}"
            orig = it.get('original') or ''
            sug = it.get('suggestion') or ''
            msg = it.get('message') or it.get('description') or ''
            ctx = self._extract_context(content, pos)

            row[0].text = str(idx)
            row[1].text = t
            row[2].text = sev
            row[3].text = pos_text

            # 原文（红色，差异片段下划线）
            self._write_diff_to_cell(row[4], orig, sug, original_side=True)
            # 建议（默认色，差异片段下划线）
            self._write_diff_to_cell(row[5], orig, sug, original_side=False)

            row[6].text = msg
            row[7].text = ctx

        # 保存并返回字节
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        with open(temp_file.name, 'rb') as f:
            content_bytes = f.read()
        os.unlink(temp_file.name)
        return content_bytes

    def _extract_context(self, content, pos):
        try:
            s = pos.get('start', 0); e = pos.get('end', 0)
            left = max(0, s - 20); right = min(len(content), e + 20)
            return content[left:right]
        except Exception:
            return ''

    def _write_diff_to_cell(self, cell, orig, sug, original_side=True):
        """将差异以 run 写入单元格。原文为红色，仅差异片段下划线；建议默认色，仅差异片段下划线。"""
        p = cell.paragraphs[0]
        # 清空默认段落内容
        for r in p.runs:
            r.clear()
        sm = difflib.SequenceMatcher(a=orig or '', b=sug or '')
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if original_side:
                seg = (orig or '')[i1:i2]
                if not seg:
                    continue
                run = p.add_run(seg)
                run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)  # 红色
                run.underline = tag in ('replace', 'delete')
            else:
                seg = (sug or '')[j1:j2]
                if not seg:
                    continue
                run = p.add_run(seg)
                # 默认颜色
                run.underline = tag in ('replace', 'insert')
        return cell

    def html_to_docx(self, html_content, title="文档", author=""):
        """
        将HTML内容转换为Word文档
        
        Args:
            html_content (str): HTML内容
            title (str): 文档标题
            author (str): 作者
        
        Returns:
            bytes: Word文档的二进制内容
        """
        # 创建新文档
        doc = Document()
        
        # 添加标题
        if title:
            doc.add_heading(title, 0)
        
        # 将HTML转换为纯文本
        text_content = self.html_converter.handle(html_content)
        
        # 处理文本内容，按段落分割
        paragraphs = text_content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # 检查是否是标题（以#开头）
                if para_text.startswith('#'):
                    # 计算标题级别
                    level = 0
                    for char in para_text:
                        if char == '#':
                            level += 1
                        else:
                            break
                    
                    # 移除#符号
                    title_text = para_text[level:].strip()
                    if title_text:
                        doc.add_heading(title_text, min(level, 9))
                else:
                    # 普通段落
                    # 处理粗体和斜体标记
                    para_text = self._process_formatting(para_text)
                    if para_text:
                        doc.add_paragraph(para_text)
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        # 读取文件内容
        with open(temp_file.name, 'rb') as f:
            content = f.read()
        
        # 删除临时文件
        os.unlink(temp_file.name)
        
        return content
    
    def _process_formatting(self, text):
        """处理文本格式标记"""
        # 移除markdown格式标记，保留纯文本
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 粗体
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # 斜体
        text = re.sub(r'`(.*?)`', r'\1', text)        # 代码
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # 链接
        
        return text.strip()
    
    def create_simple_docx(self, content, title="文档"):
        """
        创建简单的Word文档
        
        Args:
            content (str): 文档内容（纯文本）
            title (str): 文档标题
        
        Returns:
            bytes: Word文档的二进制内容
        """
        doc = Document()
        
        # 添加标题
        if title:
            doc.add_heading(title, 0)
        
        # 按段落分割内容
        paragraphs = content.split('\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        # 读取文件内容
        with open(temp_file.name, 'rb') as f:
            content = f.read()
        
        # 删除临时文件
        os.unlink(temp_file.name)
        
        return content

# 创建全局实例
document_service = DocumentService()

